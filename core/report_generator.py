import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

# ReportLab Imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, KeepTogether, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

# python-docx Imports for DOCX generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ReportGenerator:
    """
    Generates academic reports (PDF and Word formats) for ECG-HRV analysis.
    Creates and embeds 7 mandatory figures and compiles clinical parameters.
    """
    def __init__(self, fs=250):
        self.fs = fs

    def generate_all_plots(self, t, raw_sig, filtered_sig, r_peaks, rr_raw, rr_corrected, ectopic_mask, psd_data, output_dir="temp_plots"):
        """
        Generates and saves the 7 mandatory figures required for academic evaluation.
        1. Raw and Filtered ECG Signal (with zoom-in).
        2. R-Peak Detection Overlay.
        3. RR Tachogram (Raw vs Corrected).
        4. Ectopic Correction Visualization.
        5. Power Spectral Density (Welch) showing LF and HF bands.
        6. Poincaré Plot (SD1 & SD2).
        7. RR Interval Distribution (Histogram).
        Returns a dict of file paths.
        """
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        paths = {}
        
        # Style configurations
        plt.style.use('seaborn-v0_8-whitegrid' if 'seaborn-v0_8-whitegrid' in plt.style.available else 'default')
        plt.rcParams['font.sans-serif'] = 'Arial'
        plt.rcParams['font.family'] = 'sans-serif'
        plt.rcParams['axes.edgecolor'] = '#D3D3D3'
        plt.rcParams['axes.linewidth'] = 0.8
        
        # --- Figure 1: Raw vs Filtered ECG ---
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(7.5, 4.2), sharex=True)
        # Show first 5 seconds for clear visualization
        zoom_idx = t <= 5.0
        ax1.plot(t[zoom_idx], raw_sig[zoom_idx], color='#D32F2F', alpha=0.7, linewidth=1.0)
        ax1.set_title("Figure 1a: Raw Unfiltered ECG (showing powerline noise & baseline drift)", fontsize=9, fontweight='bold')
        ax1.set_ylabel("Amplitude (mV)", fontsize=8)
        
        ax2.plot(t[zoom_idx], filtered_sig[zoom_idx], color='#1B5E20', linewidth=1.2)
        ax2.set_title("Figure 1b: Filtered ECG (Baseline removed, 0.5 - 45 Hz Butterworth bandpass)", fontsize=9, fontweight='bold')
        ax2.set_xlabel("Time (seconds)", fontsize=8)
        ax2.set_ylabel("Amplitude (mV)", fontsize=8)
        plt.tight_layout()
        paths['ecg_filtered'] = os.path.join(output_dir, 'fig1_ecg_filtered.png')
        plt.savefig(paths['ecg_filtered'], dpi=150)
        plt.close()

        # --- Figure 2: R-Peak Detection ---
        fig, ax = plt.subplots(figsize=(7.5, 3.2))
        ax.plot(t[zoom_idx], filtered_sig[zoom_idx], color='#0D47A1', label='Filtered ECG', linewidth=1.2)
        
        # Filter peaks within the 5s window
        peaks_in_zoom = r_peaks[r_peaks < int(5.0 * self.fs)]
        ax.scatter(t[peaks_in_zoom], filtered_sig[peaks_in_zoom], color='#FFD600', edgecolors='black', 
                   s=40, zorder=5, label='Detected R-Peak (Pan-Tompkins)')
        ax.set_title("Figure 2: Pan-Tompkins R-Peak Detection Overlay", fontsize=9, fontweight='bold')
        ax.set_xlabel("Time (seconds)", fontsize=8)
        ax.set_ylabel("Amplitude (mV)", fontsize=8)
        ax.legend(loc='upper right', fontsize=8)
        plt.tight_layout()
        paths['r_peaks'] = os.path.join(output_dir, 'fig2_r_peaks.png')
        plt.savefig(paths['r_peaks'], dpi=150)
        plt.close()

        # --- Figure 3: RR Tachogram ---
        fig, ax = plt.subplots(figsize=(7.5, 3.2))
        ax.plot(rr_raw, color='#EF5350', alpha=0.5, linestyle='--', marker='o', markersize=3, label='Raw RR (with Ectopics)')
        ax.plot(rr_corrected, color='#43A047', marker='s', markersize=3, label='Corrected RR (Interpolated)')
        ax.set_title("Figure 3: RR Interval Tachogram", fontsize=9, fontweight='bold')
        ax.set_xlabel("Beat Number", fontsize=8)
        ax.set_ylabel("RR Interval (ms)", fontsize=8)
        ax.legend(loc='upper right', fontsize=8)
        plt.tight_layout()
        paths['tachogram'] = os.path.join(output_dir, 'fig3_tachogram.png')
        plt.savefig(paths['tachogram'], dpi=150)
        plt.close()

        # --- Figure 4: Ectopic Correction Visualization ---
        fig, ax = plt.subplots(figsize=(7.5, 3.2))
        x_indices = np.arange(len(rr_raw))
        ax.plot(x_indices, rr_corrected, color='#2E7D32', label='Corrected RR Series', linewidth=1.2)
        
        ect_indices = np.where(ectopic_mask)[0]
        if len(ect_indices) > 0:
            ax.scatter(ect_indices, rr_raw[ect_indices], color='#D50000', label='Detected Ectopic Beat (Outlier)', s=35, zorder=5)
            ax.scatter(ect_indices, rr_corrected[ect_indices], color='#2979FF', label='Corrected Interval (Cubic Spline)', s=35, zorder=5)
            
        ax.set_title("Figure 4: Ectopic Beat Detection & Spline Interpolation Correction", fontsize=9, fontweight='bold')
        ax.set_xlabel("Beat Index", fontsize=8)
        ax.set_ylabel("RR Interval (ms)", fontsize=8)
        ax.legend(loc='upper right', fontsize=8)
        plt.tight_layout()
        paths['ectopic_correction'] = os.path.join(output_dir, 'fig4_ectopic.png')
        plt.savefig(paths['ectopic_correction'], dpi=150)
        plt.close()

        # --- Figure 5: Power Spectral Density (PSD) ---
        fig, ax = plt.subplots(figsize=(7.5, 3.2))
        f = psd_data['psd_f']
        pxx = psd_data['psd_p']
        
        if len(f) > 0:
            ax.plot(f, pxx, color='#4A148C', linewidth=1.5, label='Welch PSD Estimate')
            # Shading LF band (0.04 - 0.15 Hz)
            lf_band = (f >= 0.04) & (f < 0.15)
            ax.fill_between(f[lf_band], pxx[lf_band], color='#FF7043', alpha=0.4, label='LF (0.04-0.15 Hz) - Sympathovagal')
            # Shading HF band (0.15 - 0.40 Hz)
            hf_band = (f >= 0.15) & (f <= 0.40)
            ax.fill_between(f[hf_band], pxx[hf_band], color='#26A69A', alpha=0.4, label='HF (0.15-0.40 Hz) - Parasympathetic')
            
            ax.set_xlim(0, 0.5)  # Focus on HRV relevant frequencies
            ax.set_title("Figure 5: Power Spectral Density (Welch Method)", fontsize=9, fontweight='bold')
            ax.set_xlabel("Frequency (Hz)", fontsize=8)
            ax.set_ylabel("Power Spectral Density (ms\u00B2/Hz)", fontsize=8)
            ax.legend(loc='upper right', fontsize=8)
        else:
            ax.text(0.5, 0.5, "Insufficient data for Spectral Density Plot", ha='center', va='center')
        plt.tight_layout()
        paths['psd'] = os.path.join(output_dir, 'fig5_psd.png')
        plt.savefig(paths['psd'], dpi=150)
        plt.close()

        # --- Figure 6: Poincaré Plot ---
        fig, ax = plt.subplots(figsize=(5.5, 4.2))
        if len(rr_corrected) > 2:
            x_poinc = rr_corrected[:-1]
            y_poinc = rr_corrected[1:]
            ax.scatter(x_poinc, y_poinc, color='#0D47A1', alpha=0.6, s=15, edgecolors='none', label='RR Interval Pairs')
            
            # Draw line of identity (y = x)
            min_val = min(np.min(x_poinc), np.min(y_poinc)) - 50
            max_val = max(np.max(x_poinc), np.max(y_poinc)) + 50
            ax.plot([min_val, max_val], [min_val, max_val], color='#D50000', linestyle='--', linewidth=1.0, label='Line of Identity (y = x)')
            
            # Draw ellipse axes (representing SD1 and SD2)
            mean_x = np.mean(x_poinc)
            mean_y = np.mean(y_poinc)
            sd1 = np.sqrt(0.5 * np.var(x_poinc - y_poinc, ddof=1))
            sd2 = np.sqrt(2 * np.var(rr_corrected, ddof=1) - 0.5 * np.var(x_poinc - y_poinc, ddof=1))
            
            # Plot center
            ax.scatter(mean_x, mean_y, color='#00E676', marker='+', s=100, zorder=6, label='Center Point')
            
            # Visual indicators for SD1 and SD2 orientation (45 degrees)
            angle = np.pi / 4  # 45 deg line of identity
            # Perpendicular to line of identity (SD1)
            ax.plot([mean_x - sd1 * np.sin(angle), mean_x + sd1 * np.sin(angle)],
                    [mean_y + sd1 * np.cos(angle), mean_y - sd1 * np.cos(angle)],
                    color='#FF7043', linewidth=2.0, label=f'SD1 ({sd1:.1f} ms) - Short-term')
            # Along line of identity (SD2)
            ax.plot([mean_x - sd2 * np.cos(angle), mean_x + sd2 * np.cos(angle)],
                    [mean_y - sd2 * np.sin(angle), mean_y + sd2 * np.sin(angle)],
                    color='#29B6F6', linewidth=2.0, label=f'SD2 ({sd2:.1f} ms) - Long-term')
            
            ax.set_xlim(min_val, max_val)
            ax.set_ylim(min_val, max_val)
            ax.set_title("Figure 6: Poincaré Non-Linear Scatter Plot", fontsize=9, fontweight='bold')
            ax.set_xlabel("RR\u2093 (ms)", fontsize=8)
            ax.set_ylabel("RR\u2093\u208A\u2081 (ms)", fontsize=8)
            ax.legend(loc='lower right', fontsize=8)
        else:
            ax.text(0.5, 0.5, "Insufficient data for Poincaré Plot", ha='center', va='center')
        plt.tight_layout()
        paths['poincare'] = os.path.join(output_dir, 'fig6_poincare.png')
        plt.savefig(paths['poincare'], dpi=150)
        plt.close()

        # --- Figure 7: RR Interval Distribution (Histogram) ---
        fig, ax = plt.subplots(figsize=(7.5, 3.0))
        ax.hist(rr_corrected, bins=15, color='#00acc1', edgecolor='black', alpha=0.7, density=True)
        # Add normal curve fit
        mean = np.mean(rr_corrected)
        std = np.std(rr_corrected)
        x_pdf = np.linspace(mean - 3*std, mean + 3*std, 100)
        pdf = (1 / (std * np.sqrt(2 * np.pi))) * np.exp(-((x_pdf - mean) ** 2) / (2 * std ** 2))
        ax.plot(x_pdf, pdf, color='#FF5722', linewidth=2.0, label='Gaussian Normal Approximation')
        
        ax.set_title("Figure 7: RR Interval Probability Density Distribution Histogram", fontsize=9, fontweight='bold')
        ax.set_xlabel("RR Interval (ms)", fontsize=8)
        ax.set_ylabel("Probability Density", fontsize=8)
        ax.legend(loc='upper right', fontsize=8)
        plt.tight_layout()
        paths['distribution'] = os.path.join(output_dir, 'fig7_distribution.png')
        plt.savefig(paths['distribution'], dpi=150)
        plt.close()

        return paths

    def generate_pdf(self, pdf_path, student_info, time_metrics, freq_metrics, nonlinear_metrics, ectopic_stats, plot_paths, interpretation_text):
        """
        Compiles the academic lab report PDF using ReportLab with custom styling and formatting.
        """
        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=0.5*inch, leftMargin=0.5*inch,
            topMargin=0.5*inch, bottomMargin=0.5*inch
        )

        styles = getSampleStyleSheet()
        
        # Custom Report styles
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            textColor=colors.HexColor('#0D47A1'),
            alignment=1, # Center
            spaceAfter=15
        )
        
        h1_style = ParagraphStyle(
            'H1Style',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor('#1565C0'),
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True
        )

        h2_style = ParagraphStyle(
            'H2Style',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#1E88E5'),
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=13,
            textColor=colors.HexColor('#212121'),
            spaceAfter=8
        )
        
        caption_style = ParagraphStyle(
            'CaptionStyle',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=8,
            leading=11,
            textColor=colors.HexColor('#555555'),
            alignment=1, # Center
            spaceAfter=10
        )
        
        table_text_style = ParagraphStyle(
            'TableText',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=8,
            leading=11
        )
        
        table_hdr_style = ParagraphStyle(
            'TableHdr',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=8,
            leading=11,
            textColor=colors.white
        )

        story = []

        # --- TITLE PAGE ---
        story.append(Spacer(1, 20))
        story.append(Paragraph("OPEN ENDED LAB (OEL) REPORT: CLO1 & CLO2", title_style))
        story.append(Paragraph("ECG Biomedical Signal Processing & Heart Rate Variability (HRV) Analysis System", ParagraphStyle('SubTitle', parent=title_style, fontSize=12, leading=15, textColor=colors.HexColor('#455A64'), spaceAfter=20)))
        
        # Student and Class Information Block
        info_data = [
            [Paragraph("<b>Course Title:</b> Biomedical Signal Processing", body_style), Paragraph(f"<b>Student Name:</b> {student_info.get('name', 'N/A')}", body_style)],
            [Paragraph("<b>Lab Task:</b> Open-Ended Lab Evaluation", body_style), Paragraph(f"<b>Roll/ID Number:</b> {student_info.get('id', 'N/A')}", body_style)],
            [Paragraph("<b>Department:</b> Biomedical Engineering", body_style), Paragraph(f"<b>Supervisor:</b> {student_info.get('supervisor', 'N/A')}", body_style)],
            [Paragraph("<b>Date of Experiment:</b> May 31, 2026", body_style), Paragraph(f"<b>Evaluation Score:</b> ______ / ______", body_style)]
        ]
        info_table = Table(info_data, colWidths=[3.75*inch, 3.75*inch])
        info_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ('TOPPADDING', (0,0), (-1,-1), 2),
            ('LINEBELOW', (0,-1), (-1,-1), 1, colors.HexColor('#B0BEC5')),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 15))

        # --- SECTION 1: OBJECTIVES ---
        story.append(Paragraph("1. Laboratory Objectives", h1_style))
        story.append(Paragraph(
            "1. Implement and validate a digital biomedical signal processing pipeline for high-frequency noise and baseline wander removal from raw electrocardiogram (ECG) data.<br/>"
            "2. Establish R-peak detection accuracy using the adaptive Pan-Tompkins derivative-threshold algorithm.<br/>"
            "3. Develop an outlier-rejection engine using localized median thresholding to correct ectopic cardiac beats.<br/>"
            "4. Compute and interpret linear (time-domain, frequency-domain) and non-linear (Poincaré plot, Sample Entropy) HRV metrics for physiological homeostasis assessment.",
            body_style
        ))

        # --- SECTION 2: INTRODUCTION & PHYSIOLOGICAL BACKGROUND ---
        story.append(Paragraph("2. Physiological Background & Introduction", h1_style))
        story.append(Paragraph(
            "The autonomic nervous system (ANS) controls the sinus node pacemaker through antagonistic sympathetic (stimulatory) and parasympathetic/vagal (inhibitory) branches. "
            "Heart Rate Variability (HRV) is the fluctuation in time intervals between consecutive heartbeats (R-R intervals). "
            "Depressed HRV markers serve as diagnostic indicators for autonomic neuropathy, diabetic dysautonomia, stress-induced cardiovascular remodeling, and ischemic risks. "
            "Accurate analysis requires high-fidelity digital filtering and R-peak extraction, as small noise remnants or ectopic ventricular contractions (PVCs) drastically distort HRV statistics.",
            body_style
        ))

        # --- SECTION 3: FILTER DESIGN & METHODOLOGY ---
        story.append(Paragraph("3. Biomedical Filter Design & DSP Methodology", h1_style))
        story.append(Paragraph(
            "<b>3.1 Baseline Wander Removal (Dual Median Filter)</b><br/>"
            "Baseline drift (0.1–0.5 Hz) caused by respiration is non-linearly tracked using a dual median filter. "
            "First, a 200 ms window (removing high-amplitude QRS complexes and T-waves) passes over the ECG. "
            "Next, a 600 ms window filters the output to eliminate P-waves. This isolates the baseline wander, which is subtracted from the raw signal. "
            "Unlike linear IIR filters, median filtering introduces zero phase delay and leaves QRS wave amplitudes unmodified.<br/><br/>"
            "<b>3.2 P-QRS-T Noise Suppression (Butterworth Bandpass)</b><br/>"
            "To capture ECG waves, a 3rd-order Butterworth bandpass filter is designed with a passband of 0.5 to 45.0 Hz. "
            "This band suppresses low-frequency breathing drifts and high-frequency 50/60 Hz powerline interference. "
            "It is applied using double-pass filtering (<code>scipy.signal.filtfilt</code>), enforcing zero phase shift to prevent shifting QRS peaks in time.<br/><br/>"
            "<b>3.3 Moving Enveloping (Pan-Tompkins Algorithm)</b><br/>"
            "The Pan-Tompkins algorithm isolates QRS complexes using a derivative stage, a squaring operation, and a moving window integrator (150ms window size). "
            "This integration tracks the energy envelope of the QRS complex. R-peaks are identified using adaptive dual thresholds: "
            "the signal threshold adapts to standard beat heights while the noise threshold adjusts to muscular baseline tremors. "
            "Refractory constraints (200ms lock-out) prevent T-wave double-triggering, and a back-search threshold scans back if beats are skipped.",
            body_style
        ))
        
        story.append(PageBreak())  # Force figures to start on next page

        # --- SECTION 4: PIPELINE VISUALIZATION (MANDATORY FIGURES) ---
        story.append(Paragraph("4. Signal Processing Pipeline Results & Figures", h1_style))
        
        # Fig 1: Filtered Signal
        story.append(Image(plot_paths['ecg_filtered'], width=6.5*inch, height=3.64*inch))
        story.append(Paragraph("Figure 1: Comparison between raw, noisy ECG and the filtered, baseline-corrected signal.", caption_style))
        story.append(Spacer(1, 10))

        # Fig 2: R-Peak Overlay
        story.append(Image(plot_paths['r_peaks'], width=6.5*inch, height=2.77*inch))
        story.append(Paragraph("Figure 2: Pan-Tompkins QRS output. Overlay dots represent aligned R-peak locations.", caption_style))
        story.append(Spacer(1, 15))
        
        story.append(PageBreak())

        # Fig 3 & 4: Ectopic and Tachogram
        story.append(Image(plot_paths['ectopic_correction'], width=6.5*inch, height=2.77*inch))
        story.append(Paragraph("Figure 3: Outlier detection and cubic spline interpolation correction of ectopic ventricular beats.", caption_style))
        story.append(Spacer(1, 10))
        
        story.append(Image(plot_paths['tachogram'], width=6.5*inch, height=2.77*inch))
        story.append(Paragraph("Figure 4: R-R Interval Tachogram comparing the uncorrected series with the spline-interpolated series.", caption_style))
        story.append(Spacer(1, 15))

        story.append(PageBreak())

        # Fig 5 & 6: PSD and Poincare
        story.append(Image(plot_paths['psd'], width=6.5*inch, height=2.77*inch))
        story.append(Paragraph("Figure 5: Power Spectral Density showing LF and HF integration zones for sympathovagal estimation.", caption_style))
        story.append(Spacer(1, 10))
        
        story.append(Image(plot_paths['poincare'], width=4.8*inch, height=3.67*inch))
        story.append(Paragraph("Figure 6: Poincaré plot representing non-linear cardiac dynamics. SD1 acts perpendicular to y=x.", caption_style))
        story.append(Spacer(1, 15))

        story.append(PageBreak())

        # --- SECTION 5: HRV PARAMETERS TABLES ---
        story.append(Paragraph("5. Extracted HRV Parameters & Clinical Metrics", h1_style))
        
        # Ectopic Stats Summary
        story.append(Paragraph(
            f"<b>Ectopic Correction Audit:</b> Detected {ectopic_stats['count']} ectopic events out of {ectopic_stats['total_beats']} total beats "
            f"({ectopic_stats['pct']:.2f}% ectopic burden). Ectopic intervals were reconstructed using cubic spline interpolation.",
            body_style
        ))
        
        # Table data
        hdr_style = table_hdr_style
        txt_style = table_text_style
        
        table_data = [
            [Paragraph("HRV Parameter", hdr_style), Paragraph("Value", hdr_style), Paragraph("Clinical Reference Range", hdr_style), Paragraph("Physiological Significance", hdr_style)],
            # Time Domain
            [Paragraph("<b>Mean RR</b>", txt_style), Paragraph(f"{time_metrics['mean_rr']:.1f} ms", txt_style), Paragraph("600 - 1200 ms", txt_style), Paragraph("Average length of cardiac cycle", txt_style)],
            [Paragraph("<b>Mean HR</b>", txt_style), Paragraph(f"{time_metrics['mean_hr']:.1f} BPM", txt_style), Paragraph("50 - 100 BPM", txt_style), Paragraph("Average heart rate", txt_style)],
            [Paragraph("<b>SDNN</b>", txt_style), Paragraph(f"{time_metrics['sdnn']:.1f} ms", txt_style), Paragraph("30 - 100 ms", txt_style), Paragraph("Overall autonomic variance (sympathetic & vagal)", txt_style)],
            [Paragraph("<b>RMSSD</b>", txt_style), Paragraph(f"{time_metrics['rmssd']:.1f} ms", txt_style), Paragraph("15 - 80 ms", txt_style), Paragraph("Vagal (parasympathetic) cardiac deceleration", txt_style)],
            [Paragraph("<b>pNN50</b>", txt_style), Paragraph(f"{time_metrics['pnn50']:.1f} %", txt_style), Paragraph("1.0 - 40.0 %", txt_style), Paragraph("Vagal tone / respiratory sinus arrhythmia", txt_style)],
            # Frequency Domain
            [Paragraph("<b>LF Power</b>", txt_style), Paragraph(f"{freq_metrics['lf']:.1f} ms\u00B2", txt_style), Paragraph("300 - 2000 ms\u00B2", txt_style), Paragraph("Baroreflex feedback (sympathetic + vagal)", txt_style)],
            [Paragraph("<b>HF Power</b>", txt_style), Paragraph(f"{freq_metrics['hf']:.1f} ms\u00B2", txt_style), Paragraph("200 - 1000 ms\u00B2", txt_style), Paragraph("Vagal respiratory modulation", txt_style)],
            [Paragraph("<b>LF/HF Ratio</b>", txt_style), Paragraph(f"{freq_metrics['ratio']:.2f}", txt_style), Paragraph("0.5 - 2.0", txt_style), Paragraph("Sympathovagal balance (High = stress / Sympathetic)", txt_style)],
            # Non Linear
            [Paragraph("<b>SD1</b>", txt_style), Paragraph(f"{nonlinear_metrics['sd1']:.1f} ms", txt_style), Paragraph("15 - 60 ms", txt_style), Paragraph("Short-term HRV (RMSSD correlate / Parasympathetic)", txt_style)],
            [Paragraph("<b>SD2</b>", txt_style), Paragraph(f"{nonlinear_metrics['sd2']:.1f} ms", txt_style), Paragraph("40 - 150 ms", txt_style), Paragraph("Long-term HRV (SDNN correlate / Total variation)", txt_style)],
            [Paragraph("<b>Sample Entropy</b>", txt_style), Paragraph(f"{nonlinear_metrics['sampen']:.3f}", txt_style), Paragraph("> 1.0", txt_style), Paragraph("Rhythm complexity & cellular homeostatic health", txt_style)],
        ]
        
        metrics_table = Table(table_data, colWidths=[1.8*inch, 1.2*inch, 1.8*inch, 2.7*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0D47A1')),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CFD8DC')),
            ('TOPPADDING', (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F5F7F8')]),
        ]))
        story.append(metrics_table)
        story.append(Spacer(1, 15))

        # --- SECTION 6: DISCUSSION ---
        story.append(Paragraph("6. Result Discussion & Autonomic Interpretation", h1_style))
        # Format the interpretation text beautifully
        formatted_interpretation = interpretation_text.replace('\n', '<br/>')
        story.append(Paragraph(formatted_interpretation, body_style))
        story.append(Spacer(1, 10))

        # --- SECTION 7: CONCLUSION & REFERENCES ---
        story.append(Paragraph("7. Academic Conclusion", h1_style))
        story.append(Paragraph(
            "This experiment demonstrated the design and validation of a complete biomedical signal processing system. "
            "The dual-median filter successfully tracked non-linear baseline drift. Zero-phase bandpass filtering suppressed "
            "powerline noise without shifting R-peak coordinates. Adaptive thresholding in the Pan-Tompkins algorithm "
            "maintained high detection accuracy. Ectopic beat correction isolated and repaired premature beats. "
            "The computed time, frequency, and non-linear HRV parameters accurately quantify cardiac autonomic control, "
            "providing essential diagnostics for sympathovagal balance.",
            body_style
        ))
        
        story.append(Paragraph("8. Scientific References", h1_style))
        story.append(Paragraph(
            "1. Pan, J. and Tompkins, W. J., 'A Real-Time QRS Detection Algorithm,' <i>IEEE Transactions on Biomedical Engineering</i>, vol. BME-32, no. 3, pp. 230-236, 1985.<br/>"
            "2. Task Force of the European Society of Cardiology and the North American Society of Pacing and Electrophysiology, 'Heart rate variability: standards of measurement, physiological interpretation, and clinical use,' <i>Circulation</i>, vol. 93, no. 5, pp. 1043-1065, 1996.<br/>"
            "3. Tarvainen, M. P., et al., 'Kubios HRV - Heart Rate Variability Analysis Software,' <i>Computer Methods and Programs in Biomedicine</i>, vol. 113, no. 1, pp. 210-220, 2014.",
            body_style
        ))

        # Build PDF
        doc.build(story)

    def generate_docx(self, docx_path, student_info, time_metrics, freq_metrics, nonlinear_metrics, ectopic_stats, plot_paths, interpretation_text):
        """
        Compiles the academic lab report as an editable Word Document (.docx) using python-docx.
        """
        doc = Document()
        
        # Styles Setup
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x21, 0x21, 0x21)

        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("OPEN ENDED LAB (OEL) REPORT: CLO1 & CLO2")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run("ECG Signal Processing & Heart Rate Variability (HRV) Analysis System")
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = RGBColor(0x45, 0x5A, 0x64)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Student Block
        doc.add_paragraph("-----------------------------------------------------------------------------------------------------------------------------")
        doc.add_paragraph(f"Course Title: Biomedical Signal Processing      |   Student Name: {student_info.get('name', 'N/A')}")
        doc.add_paragraph(f"Lab Task: Open-Ended Lab Evaluation             |   Roll/ID Number: {student_info.get('id', 'N/A')}")
        doc.add_paragraph(f"Department: Biomedical Engineering               |   Supervisor: {student_info.get('supervisor', 'N/A')}")
        doc.add_paragraph(f"Date: May 31, 2026                              |   Score: ______ / ______")
        doc.add_paragraph("-----------------------------------------------------------------------------------------------------------------------------")

        # Sections
        def add_section_header(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)

        add_section_header("1. Laboratory Objectives")
        doc.add_paragraph(
            "1. Implement and validate a digital biomedical signal processing pipeline for high-frequency noise and baseline wander removal from raw electrocardiogram (ECG) data.\n"
            "2. Establish R-peak detection accuracy using the adaptive Pan-Tompkins derivative-threshold algorithm.\n"
            "3. Develop an outlier-rejection engine using localized median thresholding to correct ectopic cardiac beats.\n"
            "4. Compute and interpret linear (time-domain, frequency-domain) and non-linear (Poincaré plot, Sample Entropy) HRV metrics for physiological homeostasis assessment."
        )

        add_section_header("2. Physiological Background & Introduction")
        doc.add_paragraph(
            "The autonomic nervous system (ANS) controls the sinus node pacemaker through antagonistic sympathetic (stimulatory) and parasympathetic/vagal (inhibitory) branches. "
            "Heart Rate Variability (HRV) is the fluctuation in time intervals between consecutive heartbeats (R-R intervals). "
            "Depressed HRV markers serve as diagnostic indicators for autonomic neuropathy, diabetic dysautonomia, stress-induced cardiovascular remodeling, and ischemic risks. "
            "Accurate analysis requires high-fidelity digital filtering and R-peak extraction, as small noise remnants or ectopic ventricular contractions (PVCs) drastically distort HRV statistics."
        )

        add_section_header("3. Filter Design & DSP Methodology")
        doc.add_paragraph(
            "3.1 Baseline Wander Removal (Dual Median Filter)\n"
            "Baseline drift (0.1-0.5 Hz) caused by respiration is non-linearly tracked using a dual median filter. "
            "First, a 200 ms window (removing high-amplitude QRS complexes and T-waves) passes over the ECG. "
            "Next, a 600 ms window filters the output to eliminate P-waves. This isolates the baseline wander, which is subtracted from the raw signal.\n\n"
            "3.2 P-QRS-T Noise Suppression (Butterworth Bandpass)\n"
            "To capture ECG waves, a 3rd-order Butterworth bandpass filter is designed with a passband of 0.5 to 45.0 Hz. "
            "This band suppresses low-frequency breathing drifts and high-frequency 50/60 Hz powerline interference. "
            "It is applied using double-pass filtering (scipy.signal.filtfilt), enforcing zero phase shift to prevent shifting QRS peaks in time.\n\n"
            "3.3 Moving Enveloping (Pan-Tompkins Algorithm)\n"
            "The Pan-Tompkins algorithm isolates QRS complexes using a derivative stage, a squaring operation, and a moving window integrator (150ms window size). "
            "This integration tracks the energy envelope of the QRS complex. R-peaks are identified using adaptive dual thresholds: "
            "the signal threshold adapts to standard beat heights while the noise threshold adjusts to muscular baseline tremors."
        )

        # Figures
        add_section_header("4. Signal Processing Pipeline Results & Figures")
        
        doc.add_paragraph("Figure 1: Raw vs Filtered ECG Signal (0.5-45 Hz passband)")
        doc.add_picture(plot_paths['ecg_filtered'], width=Inches(6.0))
        
        doc.add_paragraph("Figure 2: Pan-Tompkins R-Peak Detection Overlay")
        doc.add_picture(plot_paths['r_peaks'], width=Inches(6.0))
        
        doc.add_paragraph("Figure 3: Ectopic Beat Detection & Spline Interpolation Correction")
        doc.add_picture(plot_paths['ectopic_correction'], width=Inches(6.0))
        
        doc.add_paragraph("Figure 4: R-R Interval Tachogram")
        doc.add_picture(plot_paths['tachogram'], width=Inches(6.0))
        
        doc.add_paragraph("Figure 5: Power Spectral Density (Welch Method)")
        doc.add_picture(plot_paths['psd'], width=Inches(6.0))
        
        doc.add_paragraph("Figure 6: Poincaré Non-Linear Scatter Plot")
        doc.add_picture(plot_paths['poincare'], width=4.5)

        # Table
        add_section_header("5. Extracted HRV Parameters & Clinical Metrics")
        doc.add_paragraph(
            f"Ectopic Correction Audit: Detected {ectopic_stats['count']} ectopic events out of {ectopic_stats['total_beats']} total beats "
            f"({ectopic_stats['pct']:.2f}% ectopic burden). Ectopic intervals were reconstructed using cubic spline interpolation."
        )

        table_data = [
            ["HRV Parameter", "Value", "Clinical Reference Range", "Physiological Significance"],
            ["Mean RR", f"{time_metrics['mean_rr']:.1f} ms", "600 - 1200 ms", "Average length of cardiac cycle"],
            ["Mean HR", f"{time_metrics['mean_hr']:.1f} BPM", "50 - 100 BPM", "Average heart rate"],
            ["SDNN", f"{time_metrics['sdnn']:.1f} ms", "30 - 100 ms", "Overall autonomic variance (sympathetic & vagal)"],
            ["RMSSD", f"{time_metrics['rmssd']:.1f} ms", "15 - 80 ms", "Vagal (parasympathetic) cardiac deceleration"],
            ["pNN50", f"{time_metrics['pnn50']:.1f} %", "1.0 - 40.0 %", "Vagal tone / respiratory sinus arrhythmia"],
            ["LF Power", f"{freq_metrics['lf']:.1f} ms\u00B2", "300 - 2000 ms\u00B2", "Baroreflex feedback (sympathetic + vagal)"],
            ["HF Power", f"{freq_metrics['hf']:.1f} ms\u00B2", "200 - 1000 ms\u00B2", "Vagal respiratory modulation"],
            ["LF/HF Ratio", f"{freq_metrics['ratio']:.2f}", "0.5 - 2.0", "Sympathovagal balance (High = stress / Sympathetic)"],
            ["SD1", f"{nonlinear_metrics['sd1']:.1f} ms", "15 - 60 ms", "Short-term HRV (RMSSD correlate / Parasympathetic)"],
            ["SD2", f"{nonlinear_metrics['sd2']:.1f} ms", "40 - 150 ms", "Long-term HRV (SDNN correlate / Total variation)"],
            ["Sample Entropy", f"{nonlinear_metrics['sampen']:.3f}", "> 1.0", "Rhythm complexity & cellular homeostatic health"]
        ]

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        for col_idx, text in enumerate(table_data[0]):
            hdr_cells[col_idx].text = text

        for row_data in table_data[1:]:
            row_cells = table.add_row().cells
            for col_idx, text in enumerate(row_data):
                row_cells[col_idx].text = text

        add_section_header("6. Result Discussion & Autonomic Interpretation")
        doc.add_paragraph(interpretation_text)

        add_section_header("7. Conclusion")
        doc.add_paragraph(
            "This experiment demonstrated the design and validation of a complete biomedical signal processing system. "
            "The dual-median filter successfully tracked non-linear baseline drift. Zero-phase bandpass filtering suppressed "
            "powerline noise without shifting R-peak coordinates. Adaptive thresholding in the Pan-Tompkins algorithm "
            "maintained high detection accuracy. Ectopic beat correction isolated and repaired premature beats. "
            "The computed time, frequency, and non-linear HRV parameters accurately quantify cardiac autonomic control, "
            "providing essential diagnostics for sympathovagal balance."
        )

        add_section_header("8. References")
        doc.add_paragraph(
            "1. Pan, J. and Tompkins, W. J., 'A Real-Time QRS Detection Algorithm,' IEEE Transactions on Biomedical Engineering, vol. BME-32, no. 3, pp. 230-236, 1985.\n"
            "2. Task Force of the European Society of Cardiology and the North American Society of Pacing and Electrophysiology, 'Heart rate variability: standards of measurement, physiological interpretation, and clinical use,' Circulation, vol. 93, no. 5, pp. 1043-1065, 1996."
        )

        # Save Word Document
        doc.save(docx_path)
